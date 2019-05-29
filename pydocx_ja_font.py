import os
import sys

#word
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table



class pydocx_ja_font():
#日本語フォントでwordに入力を行うためのクラス
#wordのdocumentオブジェクトの生成は外で行っておくこと
#page_breakを行う場合は、一回このクラスのインスタンスを破棄(destruct)してこのクラスで入力中のパラグラフを終了させること

    def __init__(self, font_name, document_obj):#font_nameはstr型,document_objはwordのオブジェクト(参照渡し)
        self.document = document_obj
#        self.para = document_obj.add_paragraph()
        self.font = font_name

    def print(self, sentence):#sentenceはstr型で
        self.para = self.document.add_paragraph()
        run = self.para.add_run(sentence)
        run.font.name = self.font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font)

    def addTable(self,Row,Col):
        self.para = self.document.add_paragraph()
        self.table = self.document.add_table(rows=Row, cols=Col, style='Table Grid')
        #table styleの扱いについて↓tmp.docxにTable Gridのスタイルが使われていないと、python側からは使えないっぽい。要注意。
        #一度表を作成→(スタイル設定)→削除　をすればwordにTable Gridの情報が残るのでpythonから利用することができるようになる。
        #https://stackoverflow.com/questions/50687678/table-style-keyerror-uno-style-with-name-table-grid

    def typeTableCell(self,row_num,col_num,text):
        cell = self.table.cell(row_num,col_num)
        run = cell.paragraphs[0].add_run(text)
        font = run.font
        font.name = self.font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font)



document = Document('tmp.docx')#余白とかのベースを拾う

#毎回クラス作り直す page_breakをうまく動作させるため。
myWord = pydocx_ja_font("ＭＳ Ｐゴシック",document)#word_ja_fontクラスのインスタンス生成

myWord.print('こんにちは\n')
myWord.print('日本語フォントが打てます。')

myWord.addTable(2,5)
myWord.typeTableCell(0,0,"品番")
myWord.typeTableCell(0,1,"品名")
myWord.typeTableCell(0,2,"金額")
myWord.typeTableCell(0,3,"個数")
myWord.typeTableCell(0,4,"合計")


del myWord#ここでデストラクトしないとページ区切りがうまく入らない


document.add_page_break()


document.save('sample.docx')


